import os
from fpdf import FPDF
from docx import Document
from pydantic import BaseModel
from typing import List, Dict, Any, Optional

class ReportData(BaseModel):
    project_name: str
    requirements_summary: str
    # Stories array expected to have: title, description, points, priority, jira_url
    stories: List[Dict[str, Any]]
    # Sprint array expected to have: sprint_name, total_points, stories (list of titles)
    sprint_plan: List[Dict[str, Any]]

class PDFReport(FPDF):
    def header(self):
        self.set_font('Arial', 'B', 15)
        self.cell(0, 10, 'Autonomous AI Product Manager - Project Report', 0, 1, 'C')
        self.ln(5)

    def footer(self):
        self.set_y(-15)
        self.set_font('Arial', 'I', 8)
        self.cell(0, 10, f'Page {self.page_no()}', 0, 0, 'C')

def generate_pdf_report(data: ReportData, output_path: str):
    pdf = PDFReport()
    pdf.add_page()
    
    # Normalize unicode to avoid FPDF latin-1 errors
    def clean(text):
        if not text: return ""
        return str(text).encode('latin-1', 'replace').decode('latin-1')
        
    # 1. Requirements Summary
    pdf.set_font('Arial', 'B', 12)
    pdf.cell(0, 10, '1. Requirements Summary', 0, 1)
    pdf.set_font('Arial', '', 11)
    pdf.multi_cell(0, 8, clean(data.requirements_summary))
    pdf.ln(5)
    
    # 2. Sprint Plan
    pdf.set_font('Arial', 'B', 12)
    pdf.cell(0, 10, '2. Sprint Plan', 0, 1)
    for sprint in data.sprint_plan:
        pdf.set_font('Arial', 'B', 11)
        pdf.cell(0, 8, clean(f"{sprint.get('sprint_name')} ({sprint.get('total_points', 0)} pts)"), 0, 1)
        pdf.set_font('Arial', '', 10)
        for s_title in sprint.get('stories', []):
            pdf.cell(0, 6, clean(f"- {s_title}"), 0, 1)
        pdf.ln(3)
        
    # 3. User Stories & Estimation & Jira
    pdf.add_page()
    pdf.set_font('Arial', 'B', 12)
    pdf.cell(0, 10, '3. User Stories (Points & Jira Links)', 0, 1)
    
    for story in data.stories:
        pdf.set_font('Arial', 'B', 11)
        pdf.cell(0, 8, clean(f"{story.get('title')} [{story.get('priority', 'Unranked')}]"), 0, 1)
        
        pdf.set_font('Arial', '', 10)
        pdf.multi_cell(0, 6, clean(story.get('description', '')))
        
        points = story.get('points', 'Unestimated')
        pdf.cell(0, 6, clean(f"Story Points: {points}"), 0, 1)
        
        jira_url = story.get('jira_url')
        if jira_url:
            pdf.set_text_color(0, 0, 255) # Blue link
            pdf.cell(0, 6, clean(f"Jira Link: {jira_url}"), 0, 1, link=jira_url)
            pdf.set_text_color(0, 0, 0)
        
        pdf.ln(5)
        
    pdf.output(output_path)

def generate_docx_report(data: ReportData, output_path: str):
    doc = Document()
    
    doc.add_heading(f"Project Report: {data.project_name}", 0)
    
    # 1. Requirements Summary
    doc.add_heading("1. Requirements Summary", level=1)
    doc.add_paragraph(data.requirements_summary)
    
    # 2. Sprint Plan
    doc.add_heading("2. Sprint Plan", level=1)
    for sprint in data.sprint_plan:
        doc.add_heading(f"{sprint.get('sprint_name')} ({sprint.get('total_points', 0)} pts)", level=2)
        for s_title in sprint.get('stories', []):
             doc.add_paragraph(s_title, style='List Bullet')
             
    # 3. User Stories & Estimation & Jira
    doc.add_page_break()
    doc.add_heading("3. User Stories", level=1)
    
    for story in data.stories:
        doc.add_heading(f"{story.get('title')} [{story.get('priority', 'Unranked')}]", level=2)
        doc.add_paragraph(story.get('description', ''))
        
        doc.add_paragraph(f"Story Points: {story.get('points', 'Unestimated')}")
        
        jira_url = story.get('jira_url')
        if jira_url:
            # Word automatically creates clickable links for typed URLs
            doc.add_paragraph(f"Jira Link: {jira_url}")
            
    doc.save(output_path)
